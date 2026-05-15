import glob
import numpy as np

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma

from rank_bm25 import BM25Okapi

from config import embedding_model


def read_pdf(path):
    text = ""
    reader = PdfReader(path)

    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"

    return text


def read_docx(path):
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs])


def read_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_file(path):
    p = path.lower()

    if p.endswith(".pdf"):
        return read_pdf(path)

    elif p.endswith(".docx"):
        return read_docx(path)

    elif p.endswith(".txt"):
        return read_txt(path)

    return ""


text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,
    chunk_overlap=250,
    separators=[
        "\nEducation",
        "\nExperience",
        "\nSkills",
        "\nProjects",
        "\n\n",
        "\n",
        ". ",
        " ",
        ""
    ]
)


def chunk_text(text):
    return text_splitter.split_text(text)


def build_rag(resume_folder="resumes"):
    docs = []

    files = glob.glob(f"{resume_folder}/*")

    for file in files:

        lower = file.lower()

        if lower.endswith((".pdf", ".docx", ".txt")):

            text = load_file(file)

            if not text:
                continue

            chunks = chunk_text(text)

            for chunk in chunks:

                chunk = chunk.strip()

                if len(chunk) > 20:

                    docs.append(
                        Document(
                            page_content=chunk,
                            metadata={"source": file}
                        )
                    )

    vectordb = Chroma.from_documents(
        documents=docs,
        embedding=embedding_model,
        persist_directory="chroma_db"
    )

    texts = [d.page_content for d in docs]

    tokenized = [
        t.lower().split()
        for t in texts
    ]

    bm25 = BM25Okapi(tokenized)

    return vectordb, bm25, docs


def hybrid_search(
    vectordb,
    bm25,
    docs,
    job_description,
    top_k=10
):

    retriever = vectordb.as_retriever(
        search_type="mmr",
        search_kwargs={"k": top_k}
    )

    semantic_docs = retriever.invoke(job_description)

    keyword_scores = bm25.get_scores(job_description.lower().split())

    keyword_idx = np.argsort(keyword_scores)[::-1][:top_k]

    keyword_docs = [docs[i] for i in keyword_idx]

    merged = semantic_docs + keyword_docs

    unique = {}

    for d in merged:
        unique[d.metadata["source"]] = d

    return list(unique.values())[:top_k]
